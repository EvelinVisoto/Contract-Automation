"""
============================================================
Desenvolvido por: Evelin Visoto C. Fernandes
                  2024/2025
Projeto: Automação de Contratos
GitHub: https://github.com/EvelinVisoto
Descrição: Este script preenche automaticamente contratos
           com dados de uma planilha do Excel, mantendo
           a formatação do modelo original no Word.
============================================================
"""

import os
import pandas as pd
from docx import Document
from num2words import num2words

diretorio_saida = os.path.join("dist", "contratos")
os.makedirs(diretorio_saida, exist_ok=True)

# Função para substituir texto sem perder formatação
def substituir_texto_em_runs(paragrafo, substituicoes):
    for marcador, substituto in substituicoes.items():
        for run in paragrafo.runs:
            if marcador in run.text:
                run.text = run.text.replace(marcador, substituto)

# Substitui texto nos parágrafos
def preencher_paragrafos(documento, substituicoes):
    for paragrafo in documento.paragraphs:
        substituir_texto_em_runs(paragrafo, substituicoes)

# Substitui texto nas tabelas
def preencher_tabelas(documento, substituicoes):
    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_texto_em_runs(paragrafo, substituicoes)

# Verifica se todas as tags foram preenchidas
def verificar_tags(documento, substituicoes):
    for paragrafo in documento.paragraphs:
        for marcador, substituto in substituicoes.items():
            if marcador in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(marcador, substituto)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for marcador, substituto in substituicoes.items():
                        if marcador in paragrafo.text:
                            paragrafo.text = paragrafo.text.replace(marcador, substituto)

# Função que preenche o contrato
def preencher_contrato(dados, caminho_modelo, caminho_saida):
    documento = Document(caminho_modelo)

    # Ajuste dos dados
    num_endereco = str(int(dados["Número do Endereço"])) if not pd.isna(dados["Número do Endereço"]) else ""
    cnpj_mf = str(dados["CNPJ/MF"]).zfill(14).strip() if pd.notna(dados["CNPJ/MF"]) else ""

    # Formatação de valores
    valor_total_formatado = f"{dados['Valor Total']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    valor_parcela_formatado = f"{dados['Valor Parcela']:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    # Valores por extenso
    valor_total_extenso = num2words(dados["Valor Total"], lang="pt_BR")
    valor_parcela_extenso = num2words(dados["Valor Parcela"], lang="pt_BR")

    # Formatação de datas
    data_inicio_formatada = dados["Data de Início"].strftime("%d/%m/%Y") if pd.notna(dados["Data de Início"]) else ""
    data_fim_formatada = dados["Data de Fim"].strftime("%d/%m/%Y") if pd.notna(dados["Data de Fim"]) else ""

    # Dicionário de substituições
    substituicoes = {
        "{Nome}": str(dados["Nome"]),
        "{CPF}": str(dados["CPF"]),
        "{Endereco}": str(dados["Endereço"]),
        "{NumEndereco}": num_endereco,
        "{CNPJ_MF}": cnpj_mf,
        "{DataInicio}": data_inicio_formatada,
        "{DataFim}": data_fim_formatada,
        "{Duracao}": str(dados["Duração do Contrato"]),
        "{ValorTotal}": valor_total_formatado,
        "{Parcelas}": str(dados["Parcelas"]),
        "{ValorParcela}": valor_parcela_formatado,
        "{VTotal}": valor_total_extenso,
        "{VParcela}": valor_parcela_extenso,
    }

    preencher_paragrafos(documento, substituicoes)
    preencher_tabelas(documento, substituicoes)
    verificar_tags(documento, substituicoes)

    # Salva o contrato preenchido
    documento.save(caminho_saida)

# Carrega dados do Excel
dados_funcionarios = pd.read_excel("dados_funcionarios.xlsx")

# Caminho do modelo do contrato
caminho_modelo = "modelo_contrato.docx"

# Gera contratos
for indice, linha in dados_funcionarios.iterrows():
    dados = linha.to_dict()
    caminho_saida = os.path.join(diretorio_saida, f"Contrato_{dados['Nome']}.docx")
    preencher_contrato(dados, caminho_modelo, caminho_saida)
    print(f"Contrato gerado: {caminho_saida}")
